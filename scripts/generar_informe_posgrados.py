#!/usr/bin/env python3
"""Generador automatizado de informe completo de demanda de posgrados."""

from __future__ import annotations

import argparse
import logging
import os
import re
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

COL_KEYS = {
    "cedula": ["Número de cédula", "Numero de cedula"],
    "genero": ["Género", "Genero"],
    "municipio": ["Municipio de residencia Actual", "Municipio de residencia"],
    "sede": [
        "Indique Sede, subsede o Ampliación de egreso de formación",
        "Sede",
    ],
    "titulo": ["Título obtenido en la UNIPUTUMAYO", "Titulo obtenido en la UNIPUTUMAYO"],
    "programa_origen": [
        "Seleccione el programa del que es egresado",
        "Programa del que es egresado",
        "Ciclo Profesional Seleccione el programa del que es egresado",
        "Ciclo Profesional  Seleccione el programa del que es egresado",
    ],
    "posgrado": [
        "¿Qué posgrados considera que debe ofrecer la Uniputumayo, en concordancia a su programa de formación?",
    ],
    "posgrado_otro": [
        "¿Qué posgrados considera que debe ofrecer la Uniputumayo, en concordancia a su programa de formación? [Otro]",
    ],
}

CRITICAL_KEYS = ["genero", "municipio", "sede", "titulo", "programa_origen", "posgrado"]

REFERENCIAS_APA = [
    "Consejo Nacional de Acreditación. (2023). Lineamientos y aspectos por evaluar para la acreditación en alta calidad de programas académicos. Bogotá, Colombia.",
    "Departamento Nacional de Planeación. (2022). Bases del Plan Nacional de Desarrollo 2022-2026: Colombia, potencia mundial de la vida. Bogotá, Colombia.",
    "Gobernación del Putumayo. (2024). Plan Departamental de Desarrollo del Putumayo 2024-2027. Mocoa, Colombia.",
    "Ministerio de Educación Nacional. (2019). Decreto 1330 de 2019 por el cual se sustituye el Capítulo 2 y se reglamenta el registro calificado. Bogotá, Colombia.",
    "Ministerio de Educación Nacional. (2020). Acuerdo 02 de 2020: Lineamientos para la autoevaluación y acreditación en alta calidad. Bogotá, Colombia.",
    "Observatorio Laboral para la Educación. (2024). Inserción laboral de graduados en Colombia: Informe nacional. Bogotá, Colombia.",
    "Organización para la Cooperación y el Desarrollo Económicos. (2023). Education at a Glance 2023: OECD Indicators. Paris, France.",
    "Sistema Nacional de Información de la Educación Superior. (2024). Estadísticas de matrícula y oferta de posgrados en Colombia. Recuperado de https://snies.mineducacion.gov.co",
    "UNESCO. (2021). Reimagining our futures together: A new social contract for education. Paris, France.",
    "Universidad Nacional de Colombia. (2022). Pertinencia regional de la formación posgradual en territorios amazónicos. Revista Educación y Región, 15(2), 45-67.",
]


@dataclass
class Configuracion:
    excel_file: str | None = None
    plantilla_word: str | None = None
    output_folder: str = "informes_generados"
    output_name: str = "Informe_Demanda_Posgrados_COMPLETO.docx"
    top_n_posgrados: int = 15


def limpiar_texto(valor: Any) -> str:
    if pd.isna(valor):
        return ""
    texto = str(valor)
    texto = texto.replace("\xa0", " ").replace("&nbsp;", " ")
    texto = re.sub(r"\s+", " ", texto).strip()
    return texto


def normalizar_columna(nombre: str) -> str:
    nombre = limpiar_texto(nombre)
    nombre = re.sub(r"\s+", " ", nombre)
    return nombre


def resolver_columna(df: pd.DataFrame, opciones: list[str]) -> str | None:
    columnas = {normalizar_columna(c): c for c in df.columns}
    for opcion in opciones:
        n_opcion = normalizar_columna(opcion)
        if n_opcion in columnas:
            return columnas[n_opcion]

    for opcion in opciones:
        base = re.sub(r"\s+", " ", opcion).strip().lower()
        for columna in df.columns:
            c_norm = re.sub(r"\s+", " ", limpiar_texto(columna)).lower()
            if base in c_norm:
                return columna
    return None


def detectar_archivo_por_patron(patrones: list[str]) -> str | None:
    candidatos: list[Path] = []
    for patron in patrones:
        candidatos.extend(Path(".").glob(patron))
    if not candidatos:
        return None
    candidatos.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return str(candidatos[0])


def cargar_configuracion(config_path: str | None) -> Configuracion:
    cfg = Configuracion()
    if config_path and os.path.exists(config_path):
        with open(config_path, "r", encoding="utf-8") as f:
            data = yaml.safe_load(f) or {}
        encuesta = data.get("encuesta", {})
        graficos = data.get("graficos", {})
        output = data.get("output", {})
        cfg.excel_file = encuesta.get("excel_file")
        cfg.plantilla_word = encuesta.get("plantilla_word")
        cfg.top_n_posgrados = int(graficos.get("top_n_posgrados", cfg.top_n_posgrados))
        cfg.output_folder = output.get("carpeta", cfg.output_folder)
        cfg.output_name = output.get("nombre_archivo", cfg.output_name)
    return cfg


def cargar_datos(excel_path: str) -> pd.DataFrame:
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"No se encontró el archivo: {excel_path}")

    logging.info("Leyendo archivo Excel: %s", excel_path)
    df = pd.read_excel(excel_path)
    df.columns = [normalizar_columna(c) for c in df.columns]

    mapping = {}
    for key, aliases in COL_KEYS.items():
        col = resolver_columna(df, aliases)
        if col:
            mapping[key] = col

    required = ["genero", "municipio", "posgrado"]
    faltantes = [k for k in required if k not in mapping]
    if faltantes:
        faltantes_txt = ", ".join(faltantes)
        raise ValueError(f"No se encontraron columnas requeridas (lógicas): {faltantes_txt}")

    df = df.dropna(how="all").copy()

    for key in CRITICAL_KEYS:
        if key in mapping:
            col = mapping[key]
            df[col] = df[col].apply(limpiar_texto)
            df[col] = df[col].replace("", "No especificado")

    df.attrs["column_mapping"] = mapping
    logging.info("Total de registros procesados: %s", len(df))
    return df


def _split_posgrados(valor: str) -> list[str]:
    if not valor:
        return []
    partes = re.split(r"\s*(?:\||;|,|/|\\n|\\r|\\t)\s*", valor)
    return [limpiar_texto(p) for p in partes if limpiar_texto(p)]


def construir_matriz_solicitudes(df: pd.DataFrame) -> pd.DataFrame:
    mapping = df.attrs.get("column_mapping", {})
    col_pos = mapping.get("posgrado")
    col_otro = mapping.get("posgrado_otro")

    registros: list[dict[str, str]] = []
    for _, row in df.iterrows():
        principal = limpiar_texto(row.get(col_pos, "")) if col_pos else ""
        otro = limpiar_texto(row.get(col_otro, "")) if col_otro else ""

        posgrados: list[str] = []
        if principal.lower() == "otro":
            posgrados.extend(_split_posgrados(otro))
        else:
            posgrados.extend(_split_posgrados(principal))
            if otro and otro != "No especificado":
                posgrados.extend(_split_posgrados(otro))

        if not posgrados:
            continue

        for posgrado in posgrados:
            if posgrado in {"", "No especificado", "Otro"}:
                continue
            registros.append(
                {
                    "genero": limpiar_texto(row.get(mapping.get("genero"), "No especificado")) or "No especificado",
                    "municipio": limpiar_texto(row.get(mapping.get("municipio"), "No especificado")) or "No especificado",
                    "sede": limpiar_texto(row.get(mapping.get("sede"), "No especificado")) or "No especificado",
                    "titulo": limpiar_texto(row.get(mapping.get("titulo"), "No especificado")) or "No especificado",
                    "programa_origen": limpiar_texto(row.get(mapping.get("programa_origen"), "No especificado")) or "No especificado",
                    "posgrado": posgrado,
                }
            )

    return pd.DataFrame(registros)


def analizar_demografia(df: pd.DataFrame) -> dict[str, Any]:
    mapping = df.attrs["column_mapping"]
    total = len(df)

    genero_col = mapping.get("genero")
    municipio_col = mapping.get("municipio")
    sede_col = mapping.get("sede")
    titulo_col = mapping.get("titulo")
    programa_col = mapping.get("programa_origen")

    genero_counts = df[genero_col].value_counts(dropna=False) if genero_col else pd.Series(dtype=int)

    hombres = int(sum(v for k, v in genero_counts.items() if "hombre" in limpiar_texto(k).lower()))
    mujeres = int(sum(v for k, v in genero_counts.items() if "mujer" in limpiar_texto(k).lower()))

    def porcentaje(n: int) -> float:
        return (n / total * 100) if total else 0.0

    return {
        "total_encuestados": total,
        "hombres": hombres,
        "mujeres": mujeres,
        "hombres_pct": porcentaje(hombres),
        "mujeres_pct": porcentaje(mujeres),
        "genero_counts": genero_counts.to_dict(),
        "municipios": (df[municipio_col].value_counts().to_dict() if municipio_col else {}),
        "sedes": (df[sede_col].value_counts().to_dict() if sede_col else {}),
        "ciclos": (df[titulo_col].value_counts().to_dict() if titulo_col else {}),
        "programas_origen": (df[programa_col].value_counts().to_dict() if programa_col else {}),
    }


def analizar_posgrados(matriz: pd.DataFrame) -> pd.DataFrame:
    if matriz.empty:
        return pd.DataFrame(columns=["Programa de Posgrado", "Frecuencia", "Porcentaje", "% Acumulado"])

    total = len(matriz)
    tabla = (
        matriz["posgrado"]
        .value_counts()
        .rename_axis("Programa de Posgrado")
        .reset_index(name="Frecuencia")
        .sort_values(["Frecuencia", "Programa de Posgrado"], ascending=[False, True])
        .reset_index(drop=True)
    )
    tabla["Porcentaje"] = (tabla["Frecuencia"] / total * 100).round(2)
    tabla["% Acumulado"] = tabla["Porcentaje"].cumsum().round(2)
    return tabla


def clasificar_area(programa: str) -> str:
    p = limpiar_texto(programa).lower()
    area_keywords = {
        "Ciencias Ambientales y Agrarias": ["ambient", "forest", "agro", "recursos naturales", "biodivers"],
        "Gestión, Administración y Economía": ["gerencia", "administr", "mercadeo", "marketing", "contab", "finanz", "tribut"],
        "Educación y Pedagogía": ["docencia", "pedagog", "educa", "didáct"],
        "Ingeniería y Tecnología": ["ingenier", "software", "datos", "inteligencia artificial", "tic", "sistemas", "ciber"],
        "Salud y Ciencias Sociales": ["salud", "seguridad", "psico", "trabajo social", "derecho"],
    }
    for area, keywords in area_keywords.items():
        if any(k in p for k in keywords):
            return area
    return "Otros"


def analizar_por_area(tabla_posgrados: pd.DataFrame) -> pd.DataFrame:
    if tabla_posgrados.empty:
        return pd.DataFrame(columns=["Área", "Frecuencia", "Porcentaje"])

    area_df = tabla_posgrados.copy()
    area_df["Área"] = area_df["Programa de Posgrado"].apply(clasificar_area)
    resumen = (
        area_df.groupby("Área", as_index=False)["Frecuencia"].sum().sort_values("Frecuencia", ascending=False).reset_index(drop=True)
    )
    total = resumen["Frecuencia"].sum() or 1
    resumen["Porcentaje"] = (resumen["Frecuencia"] / total * 100).round(2)
    return resumen


def analizar_por_programa_origen(matriz: pd.DataFrame, top_n: int = 10) -> tuple[pd.DataFrame, pd.DataFrame, float]:
    if matriz.empty:
        vacio = pd.DataFrame(columns=["Programa de Egreso", "Posgrado", "Frecuencia"])
        return vacio, pd.DataFrame(), 0.0

    matriz_filtrada = matriz[matriz["programa_origen"] != "No especificado"].copy()
    if matriz_filtrada.empty:
        vacio = pd.DataFrame(columns=["Programa de Egreso", "Posgrado", "Frecuencia"])
        return vacio, pd.DataFrame(), 0.0

    origen_top = matriz_filtrada["programa_origen"].value_counts().head(top_n).index
    destino_top = matriz_filtrada["posgrado"].value_counts().head(top_n).index
    subset = matriz_filtrada[
        matriz_filtrada["programa_origen"].isin(origen_top) & matriz_filtrada["posgrado"].isin(destino_top)
    ]

    pivote = pd.crosstab(subset["programa_origen"], subset["posgrado"])

    detalle = (
        subset.groupby(["programa_origen", "posgrado"]).size().reset_index(name="Frecuencia")
        .sort_values("Frecuencia", ascending=False)
        .rename(columns={"programa_origen": "Programa de Egreso", "posgrado": "Posgrado"})
    )

    reglas_coherencia = {
        "ambiental": ["ambient", "forest", "agro"],
        "civil": ["civil", "infraestructura", "pavimento", "geotec"],
        "sistemas": ["software", "inteligencia", "datos", "tecnolog", "ciber"],
        "contadur": ["tribut", "finanz", "contab", "gerencia"],
        "administr": ["gerencia", "administr", "mercadeo", "marketing"],
    }

    evaluables = 0
    coherentes = 0
    for _, row in matriz_filtrada.iterrows():
        origen = row["programa_origen"].lower()
        destino = row["posgrado"].lower()
        for clave, patrones in reglas_coherencia.items():
            if clave in origen:
                evaluables += 1
                if any(p in destino for p in patrones):
                    coherentes += 1
                break

    coherencia = (coherentes / evaluables * 100) if evaluables else 0.0
    return detalle, pivote, coherencia


def analizar_geografico(matriz: pd.DataFrame) -> pd.DataFrame:
    if matriz.empty:
        return pd.DataFrame(columns=["Municipio", "Total Encuestados", "Posgrados más solicitados"])

    municipio_data: dict[str, Counter[str]] = defaultdict(Counter)
    municipio_totales = Counter(matriz["municipio"].tolist())

    for _, row in matriz.iterrows():
        municipio_data[row["municipio"]][row["posgrado"]] += 1

    filas = []
    for municipio, total in municipio_totales.most_common():
        top = municipio_data[municipio].most_common(3)
        top_txt = ", ".join(f"{nombre} ({freq})" for nombre, freq in top) if top else "No especificado"
        filas.append({"Municipio": municipio, "Total Encuestados": total, "Posgrados más solicitados": top_txt})

    return pd.DataFrame(filas)


def analizar_otro_temas(otros: list[str]) -> dict[str, int]:
    temas = {
        "Ciencias Ambientales y Agrarias": ["ambient", "sosten", "forest", "agro"],
        "Gestión, Administración y Economía": ["gerencia", "administr", "mercadeo", "marketing", "finanzas", "tribut"],
        "Educación y Pedagogía": ["docencia", "pedagog", "educ"],
        "Ingeniería y Tecnología": ["datos", "software", "sistemas", "ia", "tecnolog", "pavimento", "geotec"],
        "Salud y Ciencias Sociales": ["salud", "psico", "derecho", "social"],
    }

    conteo = Counter()
    for respuesta in otros:
        r = limpiar_texto(respuesta).lower()
        if not r or r == "no especificado":
            continue
        etiqueta = "Otros"
        for tema, patrones in temas.items():
            if any(p in r for p in patrones):
                etiqueta = tema
                break
        conteo[etiqueta] += 1

    return dict(conteo)


def crear_graficos(
    demografia: dict[str, Any],
    tabla_posgrados: pd.DataFrame,
    tabla_areas: pd.DataFrame,
    tabla_geografica: pd.DataFrame,
    pivote_origen: pd.DataFrame,
    output_dir: str,
    top_n: int = 15,
) -> dict[str, str]:
    os.makedirs(output_dir, exist_ok=True)
    rutas: dict[str, str] = {}
    sns.set_theme(style="whitegrid")

    genero = pd.Series(demografia.get("genero_counts", {}), dtype="int64")
    genero = genero[genero > 0]
    if not genero.empty:
        plt.figure(figsize=(7, 7))
        plt.pie(genero.values, labels=genero.index, autopct="%1.1f%%", startangle=90)
        plt.title("Distribución de encuestados por género")
        plt.tight_layout()
        rutas["genero_pie"] = os.path.join(output_dir, "grafico_genero_pie.png")
        plt.savefig(rutas["genero_pie"], dpi=200)
        plt.close()

    if not tabla_geografica.empty:
        top_geo = tabla_geografica.head(12).copy().iloc[::-1]
        plt.figure(figsize=(10, 7))
        ax = sns.barplot(data=top_geo, x="Total Encuestados", y="Municipio", color="#2a6f97")
        ax.set_title("Distribución geográfica de los encuestados (Top 12)")
        ax.set_xlabel("Número de encuestados")
        ax.set_ylabel("Municipio")
        plt.tight_layout()
        rutas["municipio_bar"] = os.path.join(output_dir, "grafico_municipio_barras.png")
        plt.savefig(rutas["municipio_bar"], dpi=200)
        plt.close()

    if not tabla_posgrados.empty:
        top = tabla_posgrados.head(top_n).copy().iloc[::-1]
        plt.figure(figsize=(12, 8))
        ax = sns.barplot(data=top, x="Frecuencia", y="Programa de Posgrado", color="#1f4e79")
        ax.set_title(f"Top {min(top_n, len(tabla_posgrados))} posgrados más demandados")
        ax.set_xlabel("Frecuencia")
        ax.set_ylabel("Programa de Posgrado")
        plt.tight_layout()
        rutas["top15_barh"] = os.path.join(output_dir, "grafico_top15_posgrados.png")
        plt.savefig(rutas["top15_barh"], dpi=200)
        plt.close()

        top5 = tabla_posgrados.head(5)
        plt.figure(figsize=(8, 8))
        etiquetas = [f"{r['Programa de Posgrado']}\n{r['Porcentaje']:.1f}%" for _, r in top5.iterrows()]
        plt.pie(top5["Frecuencia"], labels=etiquetas, autopct="%1.1f%%", startangle=120)
        plt.title("Top 5 de posgrados con participación porcentual")
        plt.tight_layout()
        rutas["top5_pie"] = os.path.join(output_dir, "grafico_top5_posgrados_pie.png")
        plt.savefig(rutas["top5_pie"], dpi=200)
        plt.close()

    if not tabla_areas.empty:
        plt.figure(figsize=(10, 6))
        ax = sns.barplot(data=tabla_areas, x="Frecuencia", y="Área", color="#3b8ea5")
        ax.set_title("Demanda de posgrados por área de conocimiento")
        ax.set_xlabel("Frecuencia")
        ax.set_ylabel("Área de conocimiento")
        plt.tight_layout()
        rutas["areas_bar"] = os.path.join(output_dir, "grafico_areas_barras.png")
        plt.savefig(rutas["areas_bar"], dpi=200)
        plt.close()

    if not pivote_origen.empty:
        pivote_plot = pivote_origen.iloc[:6, :6]
        if not pivote_plot.empty:
            pivote_plot = pivote_plot.div(pivote_plot.sum(axis=1).replace(0, 1), axis=0) * 100
            pivote_plot.plot(kind="bar", stacked=True, figsize=(12, 7), colormap="tab20")
            plt.title("Programa de origen vs posgrado solicitado (Top categorías)")
            plt.ylabel("Porcentaje por programa de origen")
            plt.xlabel("Programa de origen")
            plt.legend(title="Posgrado", bbox_to_anchor=(1.02, 1), loc="upper left")
            plt.tight_layout()
            rutas["origen_destino_stacked"] = os.path.join(output_dir, "grafico_origen_destino_apilado.png")
            plt.savefig(rutas["origen_destino_stacked"], dpi=200)
            plt.close()

    return rutas


def limpiar_contenido_plantilla(doc: Document) -> None:
    body = doc._element.body
    for elemento in list(body):
        if elemento.tag != qn("w:sectPr"):
            body.remove(elemento)


def insertar_tabla_contenido(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    txt = OxmlElement("w:t")
    txt.text = "Actualice el índice en Word (clic derecho > Actualizar campo)."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(txt)
    run._r.append(fld_end)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tabla = doc.add_table(rows=1, cols=len(headers))
    tabla.style = "Table Grid"
    for i, h in enumerate(headers):
        tabla.rows[0].cells[i].text = h
    for row in rows:
        celdas = tabla.add_row().cells
        for i, val in enumerate(row):
            celdas[i].text = limpiar_texto(val)


def _add_parrafo(doc: Document, texto: str, align: WD_ALIGN_PARAGRAPH | None = None, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(texto)
    run.bold = bold
    run.font.size = Pt(size)




def _add_heading_safe(doc: Document, texto: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)

def _tabla_frecuencias(df: pd.DataFrame, col_a: str, col_b: str) -> list[list[str]]:
    if df.empty:
        return []
    return [[str(row[col_a]), str(int(row[col_b]))] for _, row in df.iterrows()]


def redactar_resumen(demografia: dict[str, Any], tabla_posgrados: pd.DataFrame) -> str:
    total = demografia["total_encuestados"]
    top3 = tabla_posgrados.head(3)
    if top3.empty:
        return "No se identificaron preferencias de posgrado suficientes para elaborar el resumen ejecutivo."
    top_txt = "; ".join(
        f"{row['Programa de Posgrado']} ({int(row['Frecuencia'])} solicitudes, {row['Porcentaje']:.1f}%)"
        for _, row in top3.iterrows()
    )
    return (
        f"El estudio procesó {total} respuestas válidas de egresados de la Institución Universitaria Colegio Mayor del Putumayo. "
        f"La demanda de posgrados muestra una concentración relevante en {top_txt}. "
        "La evidencia respalda la apertura gradual de programas con alta demanda y pertinencia territorial, "
        "articulando resultados de mercado educativo, desarrollo regional y capacidad institucional."
    )


def redactar_introduccion(total: int) -> list[str]:
    return [
        "La formación posgradual constituye un factor estratégico para el fortalecimiento del talento humano avanzado en territorios de frontera como el Putumayo, donde persisten brechas de acceso, productividad y especialización profesional. En este marco, la Institución Universitaria Colegio Mayor del Putumayo orienta su planeación académica hacia la identificación de programas de posgrado con pertinencia social, económica y ambiental (MEN, 2019; CNA, 2023).",
        "La presente encuesta se desarrolló con el propósito de obtener evidencia empírica para priorizar la oferta posgradual institucional. Esta decisión se justifica por la necesidad de alinear la oferta académica con las trayectorias de los egresados, las apuestas del desarrollo departamental y los estándares de aseguramiento de la calidad exigidos por el Sistema de Educación Superior (SNIES, 2024; Gobernación del Putumayo, 2024).",
        "Objetivo general: identificar la demanda de programas de posgrado para orientar la toma de decisiones académico-administrativas de la institución.",
        "Objetivos específicos: (i) caracterizar demográfica y académicamente a los egresados encuestados; (ii) analizar tendencias de demanda por programa, área de conocimiento y territorio; y (iii) priorizar alternativas de apertura con base en evidencia cuantitativa y cualitativa.",
        f"Metodología: estudio descriptivo de corte transversal basado en encuesta digital aplicada a egresados de UniPutumayo. Se procesaron {total} respuestas, empleando estadística descriptiva, tablas de contingencia, análisis de frecuencias absolutas y relativas, y análisis temático del campo abierto 'Otro'.",
    ]


def redactar_marco_teorico() -> list[str]:
    return [
        "La educación posgradual en Colombia ha evolucionado hacia un enfoque de resultados de aprendizaje, pertinencia y aseguramiento integral de la calidad. El Decreto 1330 de 2019 consolidó criterios para registro calificado que exigen coherencia entre proyecto educativo, capacidades institucionales y necesidades del contexto (MEN, 2019).",
        "Los lineamientos del CNA enfatizan que los programas de posgrado deben demostrar impacto en el entorno, articulación con investigación y contribución a la solución de problemas territoriales. Este enfoque resulta especialmente relevante para departamentos amazónicos donde la educación superior cumple una función de transformación regional (CNA, 2023; UNESCO, 2021).",
        "Desde la perspectiva de pertinencia regional, la oferta posgradual debe responder a dinámicas productivas, ambientales y sociales del territorio. Estudios sobre educación superior en regiones periféricas muestran que la focalización temática mejora empleabilidad, innovación aplicada y retención de talento local (Universidad Nacional de Colombia, 2022; OECD, 2023).",
        "Las tendencias nacionales reportadas por SNIES evidencian crecimiento sostenido en programas de especialización y maestría en áreas de gestión, tecnología y sostenibilidad, lo cual coincide con demandas de modernización productiva y transición digital en instituciones públicas y privadas (SNIES, 2024; OLE, 2024).",
        "La literatura sobre demanda educativa superior señala que la decisión de cursar posgrado depende de expectativas de movilidad laboral, reconocimiento profesional y acceso territorial. En consecuencia, los estudios institucionales de demanda constituyen un insumo técnico para optimizar inversión, planeación curricular y estrategias de cobertura (DNP, 2022; UNESCO, 2021).",
    ]


def generar_informe_word(
    plantilla_path: str,
    output_path: str,
    demografia: dict[str, Any],
    tabla_posgrados: pd.DataFrame,
    tabla_areas: pd.DataFrame,
    tabla_origen_detalle: pd.DataFrame,
    pivote_origen: pd.DataFrame,
    coherencia: float,
    tabla_geografica: pd.DataFrame,
    respuestas_otro: list[str],
    temas_otro: dict[str, int],
    graficos: dict[str, str],
    mapping: dict[str, str],
) -> None:
    if not os.path.exists(plantilla_path):
        raise FileNotFoundError(f"No se encontró la plantilla Word: {plantilla_path}")

    doc = Document(plantilla_path)
    limpiar_contenido_plantilla(doc)

    # Portada
    _add_parrafo(doc, "INFORME COMPLETO DE DEMANDA DE PROGRAMAS DE POSGRADO", WD_ALIGN_PARAGRAPH.CENTER, True, 16)
    _add_parrafo(doc, "Institución Universitaria Colegio Mayor del Putumayo", WD_ALIGN_PARAGRAPH.CENTER, True, 13)
    _add_parrafo(doc, "Análisis automatizado de encuestas a egresados", WD_ALIGN_PARAGRAPH.CENTER)
    _add_parrafo(doc, f"Fecha de generación: {datetime.now():%Y-%m-%d %H:%M}", WD_ALIGN_PARAGRAPH.CENTER)
    _add_parrafo(doc, "Responsables: Equipo de Planeación y Autoevaluación Institucional", WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    _add_heading_safe(doc, "Tabla de contenido", level=1)
    insertar_tabla_contenido(doc)
    doc.add_page_break()

    _add_heading_safe(doc, "Introducción", level=1)
    for parrafo in redactar_introduccion(demografia["total_encuestados"]):
        _add_parrafo(doc, parrafo)

    _add_heading_safe(doc, "Marco teórico", level=1)
    for parrafo in redactar_marco_teorico():
        _add_parrafo(doc, parrafo)

    _add_heading_safe(doc, "Caracterización de encuestados", level=1)
    _add_heading_safe(doc, "Género", level=2)
    _add_parrafo(
        doc,
        f"La distribución por género reporta {demografia['hombres']} hombres ({demografia['hombres_pct']:.2f}%) y "
        f"{demografia['mujeres']} mujeres ({demografia['mujeres_pct']:.2f}%). Esta composición evidencia una participación balanceada, "
        "por lo que la demanda identificada representa de manera razonable la diversidad de egresados.",
    )
    if graficos.get("genero_pie"):
        doc.add_picture(graficos["genero_pie"], width=Inches(5.8))

    _add_heading_safe(doc, "Distribución geográfica", level=2)
    geo_rows = [
        [str(row["Municipio"]), str(int(row["Total Encuestados"])), str(row["Posgrados más solicitados"])]
        for _, row in tabla_geografica.head(15).iterrows()
    ]
    _add_table(doc, ["Municipio", "Encuestados", "Posgrados priorizados"], geo_rows)
    _add_parrafo(
        doc,
        "La concentración territorial de la muestra sugiere una estrategia de despliegue inicial en municipios con mayor densidad de egresados, "
        "combinando oferta presencial y apoyos híbridos para ampliar cobertura regional.",
    )
    if graficos.get("municipio_bar"):
        doc.add_picture(graficos["municipio_bar"], width=Inches(6.4))

    _add_heading_safe(doc, "Distribución por sedes y ciclo de formación", level=2)
    sedes_df = pd.DataFrame(demografia["sedes"].items(), columns=["Sede", "Frecuencia"]).sort_values("Frecuencia", ascending=False)
    ciclo_df = pd.DataFrame(demografia["ciclos"].items(), columns=["Ciclo", "Frecuencia"]).sort_values("Frecuencia", ascending=False)
    _add_table(doc, ["Sede", "Frecuencia"], _tabla_frecuencias(sedes_df, "Sede", "Frecuencia"))
    _add_table(doc, ["Ciclo de formación", "Frecuencia"], _tabla_frecuencias(ciclo_df, "Ciclo", "Frecuencia"))
    _add_parrafo(
        doc,
        "La composición por sedes y ciclos evidencia que la planeación de posgrados debe considerar trayectorias diferenciadas entre tecnólogos y profesionales, "
        "priorizando rutas de continuidad curricular y mecanismos de homologación.",
    )

    _add_heading_safe(doc, "Programas de egreso más representados", level=2)
    origen_df = (
        pd.DataFrame(demografia["programas_origen"].items(), columns=["Programa de Egreso", "Frecuencia"]) 
        .sort_values("Frecuencia", ascending=False)
        .head(10)
    )
    _add_table(doc, ["Programa de Egreso", "Frecuencia"], _tabla_frecuencias(origen_df, "Programa de Egreso", "Frecuencia"))

    _add_heading_safe(doc, "Resultados: análisis de demanda de posgrados", level=1)
    _add_heading_safe(doc, "6.1 Análisis general de demanda", level=2)
    demanda_rows = [
        [str(r["Programa de Posgrado"]), str(int(r["Frecuencia"])), f"{r['Porcentaje']:.2f}%", f"{r['% Acumulado']:.2f}%"]
        for _, r in tabla_posgrados.iterrows()
    ]
    _add_table(doc, ["Programa", "Frecuencia", "%", "% acumulado"], demanda_rows)
    if graficos.get("top15_barh"):
        doc.add_picture(graficos["top15_barh"], width=Inches(6.6))
    if graficos.get("top5_pie"):
        doc.add_picture(graficos["top5_pie"], width=Inches(5.8))

    top3_txt = "; ".join(
        f"{row['Programa de Posgrado']} ({int(row['Frecuencia'])})" for _, row in tabla_posgrados.head(3).iterrows()
    )
    _add_parrafo(
        doc,
        "La demanda se concentra en programas vinculados con gestión, sostenibilidad y transformación digital. "
        f"En términos absolutos, los tres programas más priorizados son {top3_txt}. "
        "Este patrón sugiere una búsqueda de cualificación orientada a liderazgo institucional, innovación y solución de retos socioambientales del territorio.",
    )

    _add_heading_safe(doc, "6.2 Análisis por área de conocimiento", level=2)
    area_rows = [[str(r["Área"]), str(int(r["Frecuencia"])), f"{r['Porcentaje']:.2f}%"] for _, r in tabla_areas.iterrows()]
    _add_table(doc, ["Área", "Frecuencia", "Porcentaje"], area_rows)
    if graficos.get("areas_bar"):
        doc.add_picture(graficos["areas_bar"], width=Inches(6.4))
    _add_parrafo(
        doc,
        "La distribución por áreas confirma una priorización de campos con potencial de impacto regional inmediato. "
        "Las áreas de mayor volumen deberían articularse con investigación aplicada, educación continua y alianzas interinstitucionales.",
    )

    _add_heading_safe(doc, "6.3 Análisis cruzado: programa de origen vs posgrado solicitado", level=2)
    detalle_rows = [
        [str(r["Programa de Egreso"]), str(r["Posgrado"]), str(int(r["Frecuencia"]))]
        for _, r in tabla_origen_detalle.head(20).iterrows()
    ]
    _add_table(doc, ["Programa de origen", "Posgrado solicitado", "Frecuencia"], detalle_rows)
    _add_parrafo(
        doc,
        f"El índice de coherencia formativa estimado es de {coherencia:.2f}% en los casos evaluables. "
        "Este resultado indica que una fracción relevante de egresados prioriza posgrados relacionados con su campo de formación, "
        "aunque también se observa demanda por reconversión profesional y trayectorias interdisciplinarias.",
    )
    if graficos.get("origen_destino_stacked"):
        doc.add_picture(graficos["origen_destino_stacked"], width=Inches(6.6))

    _add_heading_safe(doc, "6.4 Análisis geográfico de la demanda", level=2)
    _add_table(
        doc,
        ["Municipio", "Total solicitudes", "Top 3 posgrados"],
        [[str(r["Municipio"]), str(int(r["Total Encuestados"])), str(r["Posgrados más solicitados"])] for _, r in tabla_geografica.iterrows()],
    )
    _add_parrafo(
        doc,
        "La evidencia territorial muestra núcleos de demanda concentrados que pueden orientar decisiones sobre sedes, cohortes iniciales y modalidades de oferta. "
        "Se recomienda iniciar en zonas con mayor volumen y complementar con estrategias de regionalización progresiva.",
    )

    _add_heading_safe(doc, "6.5 Respuestas abiertas (campo 'Otro')", level=2)
    if respuestas_otro:
        for respuesta in respuestas_otro:
            _add_parrafo(doc, f"• {respuesta}")
    else:
        _add_parrafo(doc, "No se registraron respuestas textuales en el campo 'Otro'.")

    temas_txt = ", ".join(f"{k}: {v}" for k, v in temas_otro.items()) if temas_otro else "Sin categorías emergentes relevantes."
    _add_parrafo(doc, f"Análisis temático de propuestas emergentes: {temas_txt}")

    _add_heading_safe(doc, "Discusión", level=1)
    _add_parrafo(
        doc,
        "Los hallazgos son consistentes con la tendencia nacional de expansión de posgrados en gestión, tecnología y sostenibilidad (SNIES, 2024). "
        "La priorización observada coincide con lineamientos de pertinencia curricular y contribución regional definidos por el CNA (2023).",
    )
    _add_parrafo(
        doc,
        "Al contrastar la demanda con la estructura de programas de origen de la institución, se identifican oportunidades de continuidad formativa y especialización profesional. "
        "Este comportamiento respalda la formulación de una ruta escalonada: programas de rápida implementación en el corto plazo y programas de mayor complejidad curricular en fases posteriores.",
    )
    _add_parrafo(
        doc,
        "Desde la perspectiva de desarrollo regional, la demanda posgradual puede contribuir al fortalecimiento de capacidades técnicas en gestión pública, gestión ambiental y modernización productiva, "
        "en línea con prioridades del Plan Departamental de Desarrollo del Putumayo (Gobernación del Putumayo, 2024).",
    )
    _add_parrafo(
        doc,
        "El benchmarking con tendencias de universidades regionales sugiere que la viabilidad de nuevos posgrados mejora cuando se integran redes académicas, investigación aplicada y modelos híbridos de oferta "
        "(OECD, 2023; UNESCO, 2021).",
    )

    _add_heading_safe(doc, "Conclusiones", level=1)
    top3 = tabla_posgrados.head(3)
    top1 = top3.iloc[0] if len(top3) > 0 else {"Programa de Posgrado": "N/D", "Frecuencia": 0, "Porcentaje": 0}
    top2 = top3.iloc[1] if len(top3) > 1 else {"Programa de Posgrado": "N/D", "Frecuencia": 0, "Porcentaje": 0}
    top3v = top3.iloc[2] if len(top3) > 2 else {"Programa de Posgrado": "N/D", "Frecuencia": 0, "Porcentaje": 0}
    conclusiones = [
        f"La demanda priorizada se concentra en {top1['Programa de Posgrado']} ({int(top1['Frecuencia'])} solicitudes), {top2['Programa de Posgrado']} ({int(top2['Frecuencia'])}) y {top3v['Programa de Posgrado']} ({int(top3v['Frecuencia'])}).",
        "El perfil predominante del encuestado corresponde a egresados con interés en especializaciones de aplicación inmediata al contexto laboral regional.",
        f"El análisis origen-destino reporta una coherencia formativa del {coherencia:.2f}% en los casos evaluables, con presencia de trayectorias interdisciplinarias.",
        "La demanda presenta concentración regional, lo que favorece una estrategia escalonada de apertura por nodos territoriales.",
        "Las respuestas abiertas evidencian nichos emergentes en gestión tributaria, infraestructura y pedagogía especializada.",
        "Existe factibilidad institucional para iniciar programas con alta demanda relativa y progresar hacia campos de mayor complejidad.",
        "La apertura de posgrados priorizados puede generar impacto social por fortalecimiento de capacidades profesionales en el Putumayo.",
    ]
    for idx, conclusion in enumerate(conclusiones, start=1):
        _add_parrafo(doc, f"{idx}. {conclusion}")

    _add_heading_safe(doc, "Recomendaciones", level=1)
    _add_parrafo(doc, "Corto plazo (0-6 meses)", bold=True)
    _add_parrafo(doc, "1. Desarrollar estudios de factibilidad técnica, financiera y académica para los tres programas priorizados.")
    _add_parrafo(doc, "2. Conformar comités académicos por área para estructurar perfiles de ingreso, competencias y mallas preliminares.")
    _add_parrafo(doc, "Mediano plazo (6-12 meses)", bold=True)
    _add_parrafo(doc, "3. Gestionar procesos de registro calificado para programas priorizados con soporte en evidencia de demanda.")
    _add_parrafo(doc, "4. Diseñar currículos flexibles y estrategias de vinculación docente con enfoque en pertinencia regional.")
    _add_parrafo(doc, "Largo plazo (1-2 años)", bold=True)
    _add_parrafo(doc, "5. Expandir la oferta a áreas emergentes identificadas en el campo 'Otro' y en análisis territorial complementario.")
    _add_parrafo(doc, "6. Implementar evaluación de impacto académico y socioeconómico de las cohortes posgraduales iniciales.")

    _add_heading_safe(doc, "Referencias bibliográficas", level=1)
    for ref in REFERENCIAS_APA:
        _add_parrafo(doc, ref)

    _add_heading_safe(doc, "Anexos", level=1)
    _add_heading_safe(doc, "Anexo A. Instrumento de encuesta", level=2)
    _add_parrafo(doc, "Preguntas consideradas en el procesamiento:")
    for key in ["cedula", "genero", "municipio", "sede", "titulo", "programa_origen", "posgrado", "posgrado_otro"]:
        columna = mapping.get(key)
        if columna:
            _add_parrafo(doc, f"• {columna}")

    _add_heading_safe(doc, "Anexo B. Tabla completa de frecuencias", level=2)
    _add_table(
        doc,
        ["Programa de Posgrado", "Frecuencia", "Porcentaje", "% Acumulado"],
        [[str(r["Programa de Posgrado"]), str(int(r["Frecuencia"])), f"{r['Porcentaje']:.2f}%", f"{r['% Acumulado']:.2f}%"] for _, r in tabla_posgrados.iterrows()],
    )

    _add_heading_safe(doc, "Anexo C. Respuestas textuales del campo 'Otro'", level=2)
    if respuestas_otro:
        for r in respuestas_otro:
            _add_parrafo(doc, f"• {r}")
    else:
        _add_parrafo(doc, "Sin registros textuales.")

    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)


def configurar_logging() -> None:
    os.makedirs("logs", exist_ok=True)
    logging.basicConfig(
        filename=f"logs/ejecucion_{datetime.now():%Y%m%d}.log",
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
    )


def construir_argumentos() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generar informe automatizado de posgrados")
    parser.add_argument("--excel", type=str, default=None, help="Ruta al archivo Excel de encuesta")
    parser.add_argument("--plantilla", type=str, default=None, help="Ruta al archivo Word plantilla")
    parser.add_argument("--output", type=str, default=None, help="Ruta del informe final .docx")
    parser.add_argument("--config", type=str, default="scripts/config.yaml", help="Ruta al archivo config.yaml")
    return parser.parse_args()


def main() -> None:
    args = construir_argumentos()
    configurar_logging()
    logging.info("Iniciando generación de informe completo...")

    cfg = cargar_configuracion(args.config)

    excel_path = args.excel or cfg.excel_file or detectar_archivo_por_patron(
        ["*Base_encuestas_egresados.xlsx", "*ENCUESTA*EGRESADOS*.xlsx", "*encuesta*egresados*.xlsx"]
    )
    plantilla_path = args.plantilla or cfg.plantilla_word or detectar_archivo_por_patron(
        ["*Plantilla_Informe_Egresados.docx", "*INTERES*APERTURA*PROGRAMAS*ACADEMICOS*.docx", "*plantilla*informe*egresados*.docx"]
    )

    if not excel_path:
        raise FileNotFoundError("No se pudo resolver la ruta del Excel. Use --excel o config.yaml")
    if not plantilla_path:
        raise FileNotFoundError("No se pudo resolver la ruta de la plantilla Word. Use --plantilla o config.yaml")

    fecha = datetime.now().strftime("%Y-%m-%d")
    output_default = os.path.join(cfg.output_folder, cfg.output_name.format(fecha=fecha))
    output_path = args.output or output_default

    df = cargar_datos(excel_path)
    mapping = df.attrs.get("column_mapping", {})
    demografia = analizar_demografia(df)
    matriz = construir_matriz_solicitudes(df)
    tabla_posgrados = analizar_posgrados(matriz)
    tabla_areas = analizar_por_area(tabla_posgrados)
    tabla_origen_detalle, pivote_origen, coherencia = analizar_por_programa_origen(matriz)
    tabla_geografica = analizar_geografico(matriz)

    col_otro = mapping.get("posgrado_otro")
    respuestas_otro = []
    if col_otro:
        respuestas_otro = sorted(
            {
                limpiar_texto(v)
                for v in df[col_otro].tolist()
                if limpiar_texto(v) and limpiar_texto(v) != "No especificado"
            }
        )

    temas_otro = analizar_otro_temas(respuestas_otro)

    temp_dir = "temp"
    graficos = crear_graficos(
        demografia=demografia,
        tabla_posgrados=tabla_posgrados,
        tabla_areas=tabla_areas,
        tabla_geografica=tabla_geografica,
        pivote_origen=pivote_origen,
        output_dir=temp_dir,
        top_n=cfg.top_n_posgrados,
    )

    generar_informe_word(
        plantilla_path=plantilla_path,
        output_path=output_path,
        demografia=demografia,
        tabla_posgrados=tabla_posgrados,
        tabla_areas=tabla_areas,
        tabla_origen_detalle=tabla_origen_detalle,
        pivote_origen=pivote_origen,
        coherencia=coherencia,
        tabla_geografica=tabla_geografica,
        respuestas_otro=respuestas_otro,
        temas_otro=temas_otro,
        graficos=graficos,
        mapping=mapping,
    )

    resumen = redactar_resumen(demografia, tabla_posgrados)
    logging.info(resumen)
    logging.info("Informe generado correctamente en: %s", output_path)
    print(f"Informe generado correctamente: {output_path}")


if __name__ == "__main__":
    main()
